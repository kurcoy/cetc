import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn 
from docx.shared import RGBColor
from docx.shared import Inches
from docx.enum.text import WD_LINE_SPACING

import copy
import time
import os

from PyQt5.QtPrintSupport import  QPrinter, QPrintPreviewDialog
from PyQt5.QtGui import QTextDocument, QPixmap
from PyQt5.QtWidgets import QApplication,  QFileDialog
import tempfile
 
class myReport():
    def __init__(self):
        self.Param_result = { 'P1':0, 'P2':0, 'P':0,'p':0, 'S':0, 'S1':0, 'S2':0, 'n':0, 'N':0,'PMax':0, 'PMin':0 }
        self.Param_input  = { 'R':0,'No':0, 'h1':0, 'h2':0, 'h3':0, 'ρ1':0,'ρ2':0,'t':0, 'S0':0, 'P0':0 }
        self.type = "压力起爆装置"
        self.file = docx.Document()
        desktop = QApplication.desktop()
        self.screen_width  = desktop.screenGeometry().width()
        self.screen_height = desktop.screenGeometry().height()
        
    def GenerateWord(self):
        self.index = 0
        self.file=docx.Document()
        self.file.styles['Normal'].font.name=u'宋体'
        self.file.styles['Normal'].font.size = Pt(12)
        self.file.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        paragraph_format = self.file.styles['Normal'].paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before  = Pt(0)
        
        #Head = self.file.add_heading("",level=1)
        #Head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p = self.file.add_paragraph("")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
       #different part
        if float(self.Param_input['R']) == 0:
            self.type = "压力起爆装置"
            self.name = "YB压力起爆装置"
        else:
            self.type = "压差起爆装置"
            self.name = "YCE压差起爆装置"
            
        #run  = Head.add_run(self.type+u"-剪切销数量设计报告")
        run  = p.add_run(self.type+u"-剪切销数量设计报告")
        run.font.color.rgb = RGBColor(250,0,0)
        run.font.size = Pt(22)
        run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
              
        p = self.file.add_paragraph(u"起爆装置型号："+self.name+"    ")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(u"井号："+self.Param_input['No'])

        p = self.file.add_paragraph(u"日期：")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(str(time.strftime('%Y{y}%m{m}%d{d}').format(y='年', m='月', d='日')))
        p = self.file.add_paragraph(u" ") 
        
        if float(self.Param_input['R']) == 0:
            p = self.file.add_paragraph(str(self.index+1)+u"、确定压力起爆装置的垂直深度 h = "+self.Param_input['h1']+u"（m）")#  style='List Number')
            p = self.file.add_paragraph(u"活塞受到的压力为 P = h×ρ×g = "+self.Param_result['P']+u"（MPa）") 
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p = self.file.add_paragraph(u"ρ -- "+self.Param_input['ρ1']+u"（g/cm3）	g -- 0.0098")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            self.index = 1
        else: 
            p = self.file.add_paragraph(str(self.index+1)+u"、确定压差起爆装置的垂直深度 h1 = "+self.Param_input['h1']+u"（m）")#  style='List Number')
            p = self.file.add_paragraph(u"套管内液体产生的压力P1 = h1×ρ1×g= "+self.Param_result['P1']+u"（MPa）") 
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p = self.file.add_paragraph(u"ρ1 -- "+self.Param_input['ρ1']+u"（g/cm3）	g -- 0.0098") 
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            p = self.file.add_paragraph(str(self.index+2)+u"、确定测试阀的垂直深度 h2 = "+self.Param_input['h2']+u"（m）")#  style='List Number')
            p = self.file.add_paragraph(str(self.index+3)+u"、确定油管内液体的垂直深度 h3 = "+self.Param_input['h3']+u"（m）")#  style='List Number')
            p = self.file.add_paragraph(u"油管内液体产生的压力 P2 = h3×ρ2×g + (h1-h2)×ρ1×g = "+self.Param_result['P2']+u"（MPa）")#  style='List Number')
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p = self.file.add_paragraph(u"ρ -- "+self.Param_input['ρ1']+u"（g/cm3）	g -- 0.0098") 
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p = self.file.add_paragraph(str(self.index+4)+u"、活塞受到的压力为 P = P1-P2 = "+self.Param_result['P']+u"（MPa）") 
            self.index = 3
        #different part
        
        p = self.file.add_paragraph(str(self.index+2)+"、确定压力起爆装置所处位置温度 t = ") 
        run = p.add_run(self.Param_input['t']+u"（℃）")
    
        p = self.file.add_paragraph(str(self.index+3)+"、根据温度t查曲线，得剪切销值随温度降低百分率 δ = ") 
        run = p.add_run(self.Param_result['p']+u"%")
       
        myFile = QPixmap(":/myImages/TempChart.png")
        with tempfile.TemporaryDirectory() as tmpdirname:
            myFile.save(str(tmpdirname)+"\\TempChart.png")
            self.file.add_picture(str(tmpdirname)+"\\TempChart.png", width=Inches(4.5))
        
        p = self.file.add_paragraph(str(self.index+4)+"、查剪切销合格证，得到剪切销在常温下的剪切值 S0 = ") 
        run = p.add_run(self.Param_input['S0']+u"（MPa）")

        p = self.file.add_paragraph(str(self.index+5)+"、计算剪切销在井下的实际剪切值 S = S0×（1-δ） = ")#  style='List Number')
        run = p.add_run(self.Param_result['S']+u"（MPa）")

        p = self.file.add_paragraph(u"剪切销的最大剪切值 S1 = ") 
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(self.Param_result['S1']+u"（MPa）")

        p = self.file.add_paragraph(u"剪切销的最小剪切值 S2 = ") 
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(self.Param_result['S2']+u"（MPa）")

        p = self.file.add_paragraph(str(self.index+6)+"、剪切销的数量，安全加压 P0 = ")#  style='List Number')
        run = p.add_run(self.Param_input['P0']+u"（MPa）")
 
        p = self.file.add_paragraph(u"n = (P+P0)/S2 = ") 
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(self.Param_result['n'])
        p = self.file.add_paragraph(u"将n向上取整得N = ") 
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(self.Param_result['N'])
        run = p.add_run(u"，N为装配剪切销的数量")
 
        p = self.file.add_paragraph(str(self.index+7)+"、井口加压最大值 Pmax = (N×S1)-P = ") 
        run = p.add_run(self.Param_result['PMax']+u"（MPa）")
    
        p = self.file.add_paragraph(str(self.index+8)+"、井口加压最小值 Pmin = (N×S2)-P = ") 
        run = p.add_run(self.Param_result['PMin']+u"（MPa）")

        p = self.file.add_paragraph(u" ") 
        p = self.file.add_paragraph(u" ") 
        p = self.file.add_paragraph(u"工程师签名：                       甲方负责人：") 
        p = self.file.add_paragraph(u" ") 
        p = self.file.add_paragraph(u"年    月    日                    年    月    日      ") 
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    def Save(self):    
        #保存
        curTime = time.strftime("%Y-%m-%d_%H%M%S", time.localtime())
        file_path =  QFileDialog.getSaveFileName(None,"save file",os.getcwd()+"\\" +self.type+curTime+".docx", filter  ="docx(*.docx)|doc(*.doc)")
        if file_path[0] != "":
            self.file.save(file_path[0])
        
    def Print(self):
        #print ("print Preview")
        printer = QPrinter(QPrinter.HighResolution)
        preview = QPrintPreviewDialog(printer)
        
        preview.resize(self.screen_width,self.screen_height)
        preview.paintRequested.connect(self.handlePaintRequest)
        preview.exec_()    
    
    def handlePaintRequest(self, printer):
        document = QTextDocument( )
        document.setHtml(self.GenerateHTML())
        document.print_(printer)

    def Param_Init(self, Calc_input={}, Calc_result={} ): 
         self.Param_input  = copy.deepcopy(Calc_input)
         self.Param_result = copy.deepcopy(Calc_result)
         for key,  value in self.Param_result.items():
            self.Param_result[key] = str(value)
         for key,  value in self.Param_input.items():
            self.Param_input[key] = str(value)   
         #print(self.Param_input )

    def GenerateHTML(self):    
        self.index = 0
        if float(self.Param_input['R']) == 0:
            self.type = "压力起爆装置"
            self.name = "YB压力起爆装置"
        else:
            self.type = "压差起爆装置"
            self.name = "YCE压差起爆装置"
            
        htmldoc =""
        htmldoc +='''<!DOCTYPE html> \n<html> \n<head>\n<meta http-equiv="Content-Type" content="text/html; charset=utf-8" /> \
        \n<title>报告打印预览</title>\n<style type="text/css"> \nh1 {color:red; font-size:22pt; font-family:"宋体"} \
        \np {color:balck;font-size:12pt; font-family:"宋体";margin:5px}\n</style>\n</head> \n<body>'''#line-height:1;

        htmldoc +='''\n<h1 align="center">'''
        htmldoc +=self.type+'''-剪切销数量设计报告</h1>\n'''#<p align="left"> <br /> </p>\n'''
        htmldoc +='''<p align="center">起爆装置型号：'''+self.name +"&nbsp;&nbsp;&nbsp;井号："+self.Param_input['No']+'''</p>\n'''
        #htmldoc +='''<p align="center">井号：'''+self.Param_input['No']+"&nbsp;&nbsp;&nbsp;起爆装置型号："+self.name +'''</p>\n'''
        htmldoc +='''<p align="center">日期：'''+str(time.strftime('%Y{y}%m{m}%d{d}').format(y='年', m='月', d='日'))+'''</p>\n<p align="left"> <br /> </p>\n'''
        
        if float(self.Param_input['R']) == 0:
            htmldoc +='''<p align="left">'''+str(self.index+1)+"、确定压力起爆装置的垂直深度 h1 ="
            htmldoc +=self.Param_input['h1']+'''（m）</p>\n'''
            htmldoc +='''<p align="center">'''+"活塞受到的压力为 P = h×ρ×g = "
            htmldoc +=self.Param_result['P']+'''  （MPa）</p>\n''' 
            htmldoc +='''<p align="center">ρ1 --  '''
            htmldoc +=self.Param_input['ρ1']+''' （g/cm3）	g -- 0.0098</p>\n'''
            self.index = 0
        else:
            htmldoc +='''<p align="left">'''+str(self.index+1)+"、确定压差起爆装置的垂直深度 h1 ="
            htmldoc +=self.Param_input['h1']+'''（m）</p>\n'''
            htmldoc +='''<p align="center">套管内液体产生的压力P1 = h1×ρ1×g= '''
            htmldoc +=self.Param_result['P1']+''' （MPa）</p>\n'''
            htmldoc +='''<p align="center">ρ1 --  '''
            htmldoc +=self.Param_input['ρ1']+''' （g/cm3）	g -- 0.0098</p>\n'''
            
            htmldoc +='''<p align="left">'''+str(self.index+2)+"、确定测试阀的垂直深度 h2 = "
            htmldoc +=self.Param_input['h2']+'''（m）</p>\n'''  
            htmldoc +='''<p align="left">'''+str(self.index+3)+"、确定油管内液体的垂直深度 h3 = "
            htmldoc +=self.Param_input['h3']+'''（m）</p>\n'''  
            htmldoc +='''<p align="center">油管内液体产生的压力 P2 = h3×ρ2×g + (h1-h2)×ρ1×g = '''
            htmldoc +=self.Param_result['P2']+''' （MPa）</p>\n'''
            htmldoc +='''<p align="center">ρ2 --  '''
            htmldoc +=self.Param_input['ρ2']+''' （g/cm3）	g -- 0.0098</p>\n'''
     
            htmldoc +='''<p align="left">'''+str(self.index+4)+"、活塞受到的压力为 P = P1-P2 ="
            htmldoc +=self.Param_result['P']+'''  （MPa）</p>\n''' 
            self.index = 3
                    
        htmldoc +='''<p align="left">'''+str(self.index+2)+"、确定压差起爆装置所处位置温度 t ="
        htmldoc +=self.Param_input['t']+'''（℃）</p>\n'''  
        htmldoc +='''<p align="left">'''+str(self.index+3)+"、根据温度t查曲线，得剪切销值随温度降低百分率 δ ="
        htmldoc +=self.Param_result['p']+'''%</p>\n'''  
        
        #htmldoc += '''<p align="center"><img width=340 height=200 src="C:\\Users\\Cliff\\Desktop\\image001.png"></p>\n'''
        htmldoc += '''<p align="center"><img width=340 height=200 src=:/myImages/TempChart.png></p>\n'''
        
        htmldoc += '''<p align="left">'''+str(self.index+4)+"、查剪切销合格证，得到剪切销在常温下的剪切值 S0 ="
        htmldoc +=self.Param_input['S0']+'''（MPa）</p>\n'''  
        htmldoc += '''<p align="left">'''+str(self.index+5)+"、计算剪切销在井下的实际剪切值 S = S0×（1-δ） ="
        htmldoc +=self.Param_result['S']+'''（MPa）</p>\n'''
        htmldoc += '''<p align="center">剪切销的最大剪切值 S1 = '''
        htmldoc +=self.Param_result['S1']+'''（MPa）</p>\n'''
        htmldoc += '''<p align="center">剪切销的最小剪切值 S2 = '''
        htmldoc +=self.Param_result['S2']+'''（MPa）</p>\n'''       
        
        htmldoc += '''<p align="left">'''+str(self.index+6)+"、剪切销的数量，安全加压 P0 ="
        htmldoc +=self.Param_input['P0']+'''（MPa）</p>\n'''        
        htmldoc += '''<p align="center">n = (P+P0)/S2 ='''
        htmldoc +=self.Param_result['n']+''' </p>\n'''
        htmldoc += '''<p align="center">将n向上取整得N = '''
        htmldoc +=self.Param_result['N']+'''，N为装配剪切销的数量。</p>\n'''       

        htmldoc += '''<p align="left">'''+str(self.index+7)+"、井口加压最大值 Pmax = (N×S1)-P ="
        htmldoc +=self.Param_result['PMax']+'''（MPa）</p>\n'''         
        htmldoc += '''<p align="left">'''+str(self.index+8)+"、井口加压最小值 Pmin = (N×S2)-P ="
        htmldoc +=self.Param_result['PMin']+'''（MPa）</p>\n''' 
        
        htmldoc +='''<p align="left"> <br /> </p> \n <p align="left"> <br /> </p>\n'''
        htmldoc +='''<p align="left">工程师签名：&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;甲方负责人：</p>	\n'''
        htmldoc +='''<p align="left"> <br /> </p> \n'''
        htmldoc +='''<p align="left">&nbsp;&nbsp;&nbsp;&nbsp;年&nbsp;&nbsp;&nbsp;&nbsp;月&nbsp;&nbsp;&nbsp;&nbsp;日 \
                       &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;年&nbsp;&nbsp;&nbsp;&nbsp;月&nbsp;&nbsp;&nbsp;&nbsp;日</p>\n'''      
        htmldoc +='''</body>\n</html> '''
       
       #print(htmldoc)
        
        return htmldoc
